#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import os
import re
import pandas as pd
from docx import Document
from PyPDF2 import PdfReader
from dateutil.parser import parse
from transformers import pipeline

import subprocess


# Parent directory containing all folders
PARENT_DIR = "/Users/marcomontenegro/owl"

# Specific folders to process
FOLDERS = [
    "GWR Portfolio Submissions 20-21",
]

# Construct full paths for each folder
FOLDER_PATHS = [os.path.join(PARENT_DIR, folder) for folder in FOLDERS]

# Initialize a list to store extracted metadata
data = []
missing_files_data = []

# Updated academic year quarter date ranges
QUARTERS = {
    "Fall": [("09-17", "12-11"), ("09-20", "12-10"), ("09-19", "12-09"), ("09-21", "12-13")],
    "Winter": [("01-04", "03-19"), ("01-03", "03-18"), ("01-09", "03-24"), ("01-08", "03-22")],
    "Spring": [("03-29", "06-11"), ("03-28", "06-10"), ("03-27", "06-09"), ("03-22", "06-19")],
    "Summer": [("06-12", "09-16"), ("06-20", "09-02"), ("06-26", "09-07"), ("06-24", "09-06")],
}
VALID_YEARS = {2020, 2021,2022,2023,2024}

# Initialize Hugging Face pipelines
metacognition_analyzer = pipeline("text-classification", model="textattack/bert-base-uncased-yelp-polarity")
trait_analyzer = pipeline("text-classification", model="distilbert-base-uncased-finetuned-sst-2-english")

def extract_docx_text(filepath):
    """Extract text from a .docx file."""
    try:
        doc = Document(filepath)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    except Exception as e:
        print(f"Error reading {filepath}: {e}")
        return ""

def extract_pdf_text(filepath):
    """Extract text from a PDF file."""
    try:
        reader = PdfReader(filepath)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        print(f"Error reading {filepath}: {e}")
        return ""

def extract_doc_text(filepath):
    """Extract text from .doc files using unoconv to convert to .docx."""
    try:
        converted_path = filepath.replace(".doc", ".docx")
        subprocess.run(["unoconv", "-f", "docx", filepath], check=True)
        doc = Document(converted_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        os.remove(converted_path)
        return text
    except Exception as e:
        print(f"Error processing {filepath}: {e}")
        return ""

def extract_submission_type(filename):
    """Extract submission type (Reflection, Paper 1, Paper 2) from filename."""
    try:
        if "Reflection" in filename:
            return "Reflection"
        elif "Paper 1" in filename or "Paper1" in filename:
            return "Paper 1"
        elif "Paper 2" in filename or "Paper2" in filename:
            return "Paper 2"
        else:
            return "Unknown"
    except Exception as e:
        print(f"Error extracting submission type from filename: {e}")
        return "Unknown"

def parse_submission_date(date_str):
    """Parse a date string and return the date."""
    try:
        return parse(date_str, fuzzy=True).date()
    except Exception as e:
        print(f"Error parsing date: {date_str} -> {e}")
        return None

def assign_quarter(date):
    """Assign a quarter based on the submission date."""
    if date is None:
        return "Unknown"
    
    # Extract the month and day
    month_day = date.strftime("%m-%d")
    year = date.year

    for quarter, ranges in QUARTERS.items():
        for start, end in ranges:
            # Compare month and day for matching range
            if start <= month_day <= end:
                return f"{quarter} {year}" if year in VALID_YEARS else quarter

    return "Unknown"

def process_submission_date_and_quarter(text):
    """Extract and process submission date and assign a quarter."""
    lines = text.splitlines()
    for line in lines:
        date = parse_submission_date(line)
        if date:
            quarter = assign_quarter(date)
            if quarter != "Unknown":
                return date, quarter  # Includes quarter, with year if applicable
    return "Unknown", "Unknown"

def analyze_metacognition(text):
    """Analyze text for reflective and metacognitive content."""
    try:
        result = metacognition_analyzer(text[:512])  # Limit to 512 tokens
        return result[0]["label"], result[0]["score"]
    except Exception as e:
        print(f"Error in metacognition analysis: {e}")
        return "Error", 0.0

def analyze_traits(text):
    """Analyze text for the five traits."""
    traits = ["Purpose", "Synthesis", "Support", "Style", "Mechanics"]
    trait_scores = {trait: 0.0 for trait in traits}
    try:
        for trait in traits:
            result = trait_analyzer(f"{trait}: {text[:512]}")  # Customize for each trait
            trait_scores[trait] = result[0]["score"]
    except Exception as e:
        print(f"Error analyzing traits: {e}")
    return trait_scores

def categorize_paper(text, submission_type):
    """Categorize the type of paper based on its content and submission type."""
    if submission_type == "Reflection":
        return "Reflection"

    text_lower = text.lower()

     # Define keyword groups for categorization
    categories = {
        "Lab Report": [
            "experiment", "results", "methodology", "materials and methods",
            "observations", "scientific study", "data analysis", "hypothesis",
            "laboratory", "practical work", "technical report", "variables"
        ],
        "Essay": [
            "thesis", "argument", "critical analysis", "analysis of",
            "persuasive", "compare and contrast", "expository", "point of view",
            "rhetorical", "philosophical", "analytical essay", "interpretation"
        ],
        "Research Paper": [
            "data", "research", "empirical study", "statistical analysis",
            "field study", "hypothesis", "literature review", "citation",
            "bibliography", "quantitative", "qualitative", "case study"
        ],
        "Creative Writing": [
            "creative writing", "poem", "short story", "narrative", "fiction",
            "prose", "imaginative", "memoir", "autobiographical", "artistic"
        ],
        "Case Study": [
            "case study", "business case", "scenario analysis", "case analysis",
            "problem solving", "diagnosis", "situational analysis", "solution"
        ],
        "Presentation": [
            "presentation", "slides", "powerpoint", "visuals", "graphics",
            "multimedia", "bullet points", "oral presentation", "poster"
        ],
        "Report": [
            "report", "summary", "overview", "findings", "executive summary",
            "progress report", "white paper", "informational document"
        ]
    }

    # Iterate through categories to find a match
    for category, keywords in categories.items():
        if any(keyword in text_lower for keyword in keywords):
            return category

    # Fallback categorization based on the submission type
    if "reflection" in text_lower or "self-assessment" in text_lower:
        return "Reflection"
    if "lab" in text_lower or "experiment" in text_lower:
        return "Lab Report"
    if "analysis" in text_lower or "essay" in text_lower:
        return "Essay"
    if "data" in text_lower or "research" in text_lower:
        return "Research Paper"
    if "creative" in text_lower or "story" in text_lower:
        return "Creative Writing"
    if "case" in text_lower or "study" in text_lower:
        return "Case Study"
    if "presentation" in text_lower or "slides" in text_lower:
        return "Presentation"
    if "report" in text_lower:
        return "Report"

    return "Other"

def extract_name_from_filename(filename):
    """Extract name from the filename."""
    try:
        match = re.search(r"_(\D+?)_([A-Z])_", filename)
        if match:
            last_name = match.group(1).capitalize()
            first_initial = match.group(2).capitalize()
            return f"{first_initial}. {last_name}"
        return "Unknown"
    except Exception as e:
        print(f"Error extracting name from filename: {e}")
        return "Unknown"

def parse_metadata_from_text(text, filename):
    """Extract metadata and classify paper."""
    try:
        name = extract_name_from_filename(filename)
        submission_type = extract_submission_type(filename)
        submission_date, quarter = process_submission_date_and_quarter(text)
        paper_type = categorize_paper(text, submission_type)
        reflective_label, reflective_score = analyze_metacognition(text)
        trait_scores = analyze_traits(text)

        return {
            "name": name,
            "paper_type": paper_type,
            "submission_type": submission_type,
            "submission_date": submission_date.strftime("%Y-%m-%d") if submission_date != "Unknown" else "Unknown",
            "quarter": quarter,
            "word_count": len(text.split()),
            "metacognition_label": reflective_label,
            "metacognition_score": reflective_score,
            **trait_scores
        }
    except Exception as e:
        print(f"Error parsing metadata for {filename}: {e}")
        return {
            "name": "Unknown",
            "paper_type": "Unknown",
            "submission_type": "Unknown",
            "submission_date": "Unknown",
            "quarter": "Unknown",
            "word_count": 0,
            "metacognition_label": "Error",
            "metacognition_score": 0.0,
            "Purpose": 0.0,
            "Synthesis": 0.0,
            "Support": 0.0,
            "Style": 0.0,
            "Mechanics": 0.0,
        }

def process_files(folder_path):
    """Process files in the specified folder."""
    for root, _, files in os.walk(folder_path):
        for file in files:
            filepath = os.path.join(root, file)
            if file.endswith(".docx"):
                text = extract_docx_text(filepath)
            elif file.endswith(".pdf"):
                text = extract_pdf_text(filepath)
            elif file.endswith(".doc"):
                text = extract_doc_text(filepath)
            else:
                print(f"Unsupported file type: {file}")
                missing_files_data.append({
                    "filename": file,
                    "filepath": filepath,
                    "name": "Unknown",
                    "paper_type": "Unknown",
                    "submission_type": "Unknown",
                    "submission_date": "Unknown",
                    "quarter": "Unknown",
                    "word_count": 0,
                    "metacognition_label": "Unknown",
                    "metacognition_score": 0.0,
                    "Purpose": 0.0,
                    "Synthesis": 0.0,
                    "Support": 0.0,
                    "Style": 0.0,
                    "Mechanics": 0.0,
                })
                continue

            metadata = parse_metadata_from_text(text, file)
            metadata["filename"] = file
            metadata["filepath"] = filepath
            data.append(metadata)

# Process files in all specified folders
if __name__ == "__main__":
    for folder_path in FOLDER_PATHS:
        process_files(folder_path)

# Save metadata to a DataFrame
df = pd.DataFrame(data)

# Separate DataFrames for each type
reflection_df = df[df["submission_type"] == "Reflection"]
paper1_df = df[df["submission_type"] == "Paper 1"]
paper2_df = df[df["submission_type"] == "Paper 2"]
missing_files_df = pd.DataFrame(missing_files_data)

# Export DataFrames
reflection_df.to_csv("reflection_submissions.csv", index=False)
paper1_df.to_csv("paper1_submissions.csv", index=False)
paper2_df.to_csv("paper2_submissions.csv", index=False)
missing_files_df.to_csv("missing_files.csv", index=False)

print("Submissions categorized and saved to separate files: reflection, paper1, paper2, and missing files.")
